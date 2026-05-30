import importlib
import sys
import types
import unittest
from pathlib import Path

from docx import Document

from main import (
    build_marketplace_search_links,
    load_candidate_links_file,
    load_candidate_products_file,
    run_stage14_v2a_delivery,
)


ROOT = Path(__file__).resolve().parents[1]
DEMO_WORD = ROOT / "data" / "办公屏风01模板样例.docx"
DEMO_LINKS = ROOT / "data" / "stage14_candidate_links_demo.txt"
DEMO_CSV = ROOT / "data" / "stage14_candidate_products_demo.csv"


class _StreamlitTrap(types.ModuleType):
    def __init__(self) -> None:
        super().__init__("streamlit")
        self.accessed: list[str] = []

    def __getattr__(self, name: str) -> object:
        self.accessed.append(name)
        raise AssertionError(f"streamlit.{name} should not be accessed during import")


def _workspace(name: str) -> Path:
    path = ROOT / "outputs" / "stage14_v2a_test_workspace" / name
    path.mkdir(parents=True, exist_ok=True)
    return path


def _extract_docx_text(path: Path) -> str:
    document = Document(str(path))
    parts: list[str] = []
    parts.extend(paragraph.text for paragraph in document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


class Stage14V2ASourcingFlowTests(unittest.TestCase):
    def test_search_links_explain_browser_assisted_boundary(self) -> None:
        search = build_marketplace_search_links(
            [{"name": "办公屏风", "required_value": "1600x750x750mm 钢制框架 支持现场安装"}]
        )

        self.assertIn("办公屏风", search["taobao_keyword"])
        self.assertIn("1600x750x750", search["jd_keyword"])
        self.assertIn("s.taobao.com/search", search["taobao_search_url"])
        self.assertIn("search.jd.com/Search", search["jd_search_url"])
        self.assertIn("浏览器", search["notice"])
        self.assertIn("不强抓", search["notice"])
        self.assertIn("Cookie", search["notice"])

    def test_candidate_links_demo_can_be_standardized_without_fetching_pages(self) -> None:
        products = load_candidate_links_file(DEMO_LINKS)

        self.assertEqual(len(products), 4)
        self.assertEqual(products[0]["adapter_source_type"], "manual_link_list")
        self.assertEqual(products[0]["platform"], "淘宝")
        self.assertEqual(products[0]["price"], 998)
        self.assertIn("1600x750x750", products[0]["dimensions"])
        self.assertIn("钢制", products[0]["material"])
        self.assertIn("安装", products[0]["installation_service"])
        self.assertEqual(products[1]["platform"], "京东")
        self.assertEqual(products[2]["platform"], "天猫")

    def test_plain_link_is_kept_as_pending_manual_review_not_fabricated(self) -> None:
        products = load_candidate_links_file(DEMO_LINKS)
        plain_link = products[-1]

        self.assertTrue(plain_link["title"].startswith("待人工复核候选商品"))
        self.assertIn("title", plain_link["missing_fields"])
        self.assertIn("price", plain_link["missing_fields"])
        self.assertIn("dimensions", plain_link["missing_fields"])
        self.assertIn("material", plain_link["missing_fields"])
        self.assertTrue(plain_link["manual_review_required"])
        self.assertEqual(plain_link["price"], None)
        self.assertEqual(plain_link["material"], "")

    def test_candidate_products_csv_demo_still_enters_same_flow(self) -> None:
        products = load_candidate_products_file(DEMO_CSV)

        self.assertEqual(len(products), 3)
        self.assertEqual(products[0]["platform"], "淘宝")
        self.assertEqual(products[1]["platform"], "京东")
        self.assertEqual(products[2]["platform"], "天猫")
        self.assertIn("现场安装", products[0]["installation_service"])

    def test_stage14_v2a_delivery_generates_word_from_link_list(self) -> None:
        workspace = _workspace("delivery_from_links")
        result = run_stage14_v2a_delivery(DEMO_WORD, DEMO_LINKS, output_dir=workspace)

        self.assertGreater(result["requirements_count"], 0)
        self.assertEqual(result["products_count"], 4)
        self.assertGreaterEqual(result["top_products_count"], 3)
        self.assertTrue(Path(result["output_docx_path"]).exists())
        self.assertTrue(Path(result["imported_candidate_products_path"]).exists())
        self.assertTrue(Path(result["ranked_products_path"]).exists())
        self.assertIn("top_products", result)

        text = _extract_docx_text(Path(result["output_docx_path"]))
        self.assertIn("淘宝/京东寻源辅助", text)
        self.assertIn("候选商品 Top 3", text)
        self.assertIn("待复核字段", text)
        self.assertIn("不强抓", text)

    def test_ui_streamlit_imports_and_exposes_v2a_action(self) -> None:
        previous_streamlit = sys.modules.get("streamlit")
        trap = _StreamlitTrap()
        sys.modules["streamlit"] = trap
        sys.modules.pop("ui_streamlit", None)

        try:
            module = importlib.import_module("ui_streamlit")
        finally:
            if previous_streamlit is None:
                sys.modules.pop("streamlit", None)
            else:
                sys.modules["streamlit"] = previous_streamlit

        self.assertTrue(callable(getattr(module, "render_app", None)))
        self.assertTrue(callable(getattr(module, "run_stage14_v2a_delivery_action", None)))
        self.assertEqual(trap.accessed, [])

    def test_no_disallowed_implementation_markers(self) -> None:
        checked_files = [ROOT / "main.py", ROOT / "ui_streamlit.py", ROOT / "src" / "doc_writer.py"]
        combined_text = "\n".join(path.read_text(encoding="utf-8") for path in checked_files)

        for phrase in ["本系统会保存账号密码", "偷 Cookie", "使用代理池", "绕过验证码", "可直接下单", "本报告为最终采购结论"]:
            self.assertNotIn(phrase, combined_text)


if __name__ == "__main__":
    unittest.main()
