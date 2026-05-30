import unittest
from pathlib import Path

from docx import Document

from src.doc_parser import REQUIRED_HEADERS
from src.doc_writer import write_responses


def _create_input_docx(path: Path) -> None:
    document = Document()
    table = document.add_table(rows=2, cols=len(REQUIRED_HEADERS))
    for index, header in enumerate(REQUIRED_HEADERS):
        table.rows[0].cells[index].text = header

    values = ["1", "屏风高度", "mm", "不低于 1200", ""]
    for index, value in enumerate(values):
        table.rows[1].cells[index].text = value

    document.save(str(path))


def _extract_docx_text(path: Path) -> str:
    document = Document(str(path))
    parts: list[str] = []
    parts.extend(paragraph.text for paragraph in document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def _workspace(name: str) -> Path:
    output_dir = Path(__file__).resolve().parents[1] / "outputs" / "stage12_5_test_workspace" / name
    output_dir.mkdir(parents=True, exist_ok=True)
    return output_dir


class Stage125WordDeliveryReportTests(unittest.TestCase):
    def test_write_responses_legacy_three_arg_call_still_works(self) -> None:
        workspace = _workspace("legacy_call")
        input_docx = workspace / "input.docx"
        output_docx = workspace / "output.docx"
        _create_input_docx(input_docx)

        result = write_responses(
            input_docx,
            [{"index": "1", "response_value": "满足，需结合候选商品人工复核。"}],
            output_docx,
        )

        self.assertEqual(result, output_docx)
        self.assertTrue(output_docx.exists())
        text = _extract_docx_text(output_docx)
        self.assertIn("满足，需结合候选商品人工复核。", text)

    def test_without_delivery_metadata_old_word_output_is_generated(self) -> None:
        workspace = _workspace("old_output")
        input_docx = workspace / "input.docx"
        output_docx = workspace / "output.docx"
        _create_input_docx(input_docx)

        write_responses(
            input_docx,
            [{"index": "1", "response_value": "候选响应值"}],
            output_docx,
            summary_products=[{"title": "候选屏风", "score": 10}],
        )

        text = _extract_docx_text(output_docx)
        self.assertIn("候选商品 Top 3", text)
        self.assertIn("候选屏风", text)

    def test_delivery_metadata_adds_client_readable_sections(self) -> None:
        workspace = _workspace("delivery_metadata")
        input_docx = workspace / "input.docx"
        output_docx = workspace / "output.docx"
        _create_input_docx(input_docx)

        write_responses(
            input_docx,
            [{"index": "1", "response_value": "候选响应值"}],
            output_docx,
            delivery_metadata={
                "data_source_priority": [
                    "reviewed_products.json",
                    "crawler_products.json",
                    "sample_products.json",
                ],
                "actual_data_source": "sample_products.json（本地样例商品兜底）",
                "products_source_path": "data/sample_products.json",
                "requirements_count": 1,
                "products_count": 3,
                "ranked_products_count": 3,
                "top_products_count": 3,
                "output_paths": {
                    "recommendation_result_docx": "outputs/recommendation_result.docx",
                    "ranked_products_json": "outputs/ranked_products.json",
                    "responses_json": "outputs/responses.json",
                },
                "review_status": {
                    "has_reviewed_products": True,
                    "has_review_report": True,
                    "reviewed_products_count": 1,
                    "approved_count": 1,
                    "rejected_count": 0,
                    "needs_more_info_count": 0,
                    "unmatched_review_items_count": 0,
                    "risk_note": "仍需人工复核候选商品和关键字段。",
                },
            },
        )

        text = _extract_docx_text(output_docx)
        self.assertIn("交付报告摘要", text)
        self.assertIn("采购辅助候选", text)
        self.assertIn("不代表最终采购结论", text)
        self.assertIn("数据来源说明", text)
        self.assertIn("人工复核状态说明", text)
        self.assertIn("字段风险提示", text)
        self.assertIn("输出文件索引", text)
        self.assertNotIn("本报告为最终采购结论", text)
        self.assertNotIn("本报告是最终采购结论", text)
        self.assertNotIn("作为最终采购结论", text)


if __name__ == "__main__":
    unittest.main()
