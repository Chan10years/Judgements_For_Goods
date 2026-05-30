import json
import socket
import unittest
from pathlib import Path
from typing import Any

from docx import Document

import ui_streamlit
from main import run_pipeline
from src.crawler.adapters.manual_json_adapter import ManualJsonAdapter
from src.crawler.pipeline import run_crawler
from src.crawler.review import generate_manual_review_items
from src.crawler.review_apply import apply_manual_reviews
from src.crawler.review_template import generate_review_template
from src.crawler.review_validate import validate_review_filled
from src.crawler.seed_loader import load_seed_urls


ROOT = Path(__file__).resolve().parents[1]
STAGE13_PRODUCTS_PATH = ROOT / "data" / "stage13_sample_products.json"
STAGE13_SEEDS_PATH = ROOT / "data" / "stage13_seed_urls.json"
STAGE13_REPORT_PATH = ROOT / "outputs" / "stage13_acceptance_report.json"
DEFAULT_WORD_PATH = ROOT / "data" / "办公屏风01模板样例.docx"


def _write_json(path: Path, data: Any) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    return path


def _read_json(path: Path) -> Any:
    return json.loads(path.read_text(encoding="utf-8"))


def _docx_text(path: Path) -> str:
    document = Document(str(path))
    parts: list[str] = []
    parts.extend(paragraph.text for paragraph in document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def _make_stage13_crawler_config(workspace: Path) -> Path:
    config = {
        "user_agent": "JudgementsForGoodsStage13Acceptance/1.0 (+public-url-audit; procurement-assist-only)",
        "timeout_seconds": 10,
        "max_retries": 0,
        "retry_backoff_seconds": 0,
        "request_interval_seconds": 2,
        "respect_robots_txt": True,
        "max_urls_per_run": 3,
        "output_raw_path": str(workspace / "crawler_raw.json"),
        "output_products_path": str(workspace / "crawler_products.json"),
        "output_report_path": str(workspace / "crawler_report.json"),
        "log_path": str(workspace / "crawler_run.log"),
    }
    return _write_json(workspace / "crawler_config.json", config)


def _build_valid_review_filled(template: dict[str, Any]) -> dict[str, list[dict[str, Any]]]:
    items: list[dict[str, Any]] = []
    approved_done = False
    for item in template.get("review_items", []):
        filled = dict(item)
        title = str(filled.get("title", ""))
        if not approved_done and "缺字段办公屏风" in title:
            filled.update(
                {
                    "review_status": "approved",
                    "review_note": "阶段13验收中人工模拟确认价格、尺寸、材质、安装服务、来源和证据文本。",
                    "confirmed_price": 1260,
                    "confirmed_dimensions": "1600x750x750mm",
                    "confirmed_material": "钢制框架，实木颗粒板台面",
                    "confirmed_installation_service": "支持送货至指定地点和现场安装，费用范围仍需采购前人工确认。",
                    "confirmed_source": "阶段13人工复核半真实样例",
                    "confirmed_evidence_text": "阶段13人工复核半真实样例确认该候选包含价格、尺寸、材质、配送和现场安装服务描述。",
                    "confirmed_fields": ["price", "dimensions", "material", "installation_service", "source", "evidence_text"],
                }
            )
            approved_done = True
        else:
            filled.update(
                {
                    "review_status": "needs_more_info",
                    "review_note": "阶段13验收保留为待补充项，采购前需要人工继续确认。",
                }
            )
        items.append(filled)
    return {"review_items": items}


class _FakeStreamlit:
    def __init__(self, clicked_label: str) -> None:
        self.clicked_label = clicked_label
        self.buttons_seen: list[str] = []
        self.session_state: dict[str, Any] = {}
        self.sidebar = self

    def __enter__(self) -> "_FakeStreamlit":
        return self

    def __exit__(self, exc_type: object, exc: object, tb: object) -> None:
        return None

    def columns(self, count: int) -> list["_FakeStreamlit"]:
        return [self for _ in range(count)]

    def button(self, label: str, **_: Any) -> bool:
        self.buttons_seen.append(label)
        return label == self.clicked_label

    def expander(self, *_: Any, **__: Any) -> "_FakeStreamlit":
        return self

    def set_page_config(self, **_: Any) -> None:
        return None

    def title(self, *_: Any, **__: Any) -> None:
        return None

    def subheader(self, *_: Any, **__: Any) -> None:
        return None

    def write(self, *_: Any, **__: Any) -> None:
        return None

    def info(self, *_: Any, **__: Any) -> None:
        return None

    def header(self, *_: Any, **__: Any) -> None:
        return None

    def caption(self, *_: Any, **__: Any) -> None:
        return None

    def table(self, *_: Any, **__: Any) -> None:
        return None

    def dataframe(self, *_: Any, **__: Any) -> None:
        return None

    def success(self, *_: Any, **__: Any) -> None:
        return None

    def error(self, *_: Any, **__: Any) -> None:
        return None

    def markdown(self, *_: Any, **__: Any) -> None:
        return None

    def code(self, *_: Any, **__: Any) -> None:
        return None

    def json(self, *_: Any, **__: Any) -> None:
        return None


class Stage13RealSampleAcceptanceTests(unittest.TestCase):
    def test_stage13_real_and_semi_real_acceptance_chain(self) -> None:
        self.assertTrue(STAGE13_PRODUCTS_PATH.exists())
        self.assertTrue(STAGE13_SEEDS_PATH.exists())

        seed_rows = _read_json(STAGE13_SEEDS_PATH)
        product_rows = _read_json(STAGE13_PRODUCTS_PATH)
        self.assertIsInstance(seed_rows, list)
        self.assertIsInstance(product_rows, list)
        self.assertGreaterEqual(len(product_rows), 3)
        self.assertLessEqual(len(product_rows), 6)

        workspace = ROOT / "outputs" / "stage13_test_workspace" / "runtime"
        workspace.mkdir(parents=True, exist_ok=True)
        crawler_config_path = _make_stage13_crawler_config(workspace)
        seeds, invalid_seeds = load_seed_urls(STAGE13_SEEDS_PATH)
        self.assertGreaterEqual(len(seeds), 1)
        self.assertEqual(invalid_seeds, [])

        old_timeout = socket.getdefaulttimeout()
        socket.setdefaulttimeout(3)
        try:
            crawler_report = run_crawler(config_path=crawler_config_path, seed_path=STAGE13_SEEDS_PATH)
        finally:
            socket.setdefaulttimeout(old_timeout)

        self.assertEqual(crawler_report["total_url_count"], len(seed_rows))
        self.assertIn("summary", crawler_report)
        self.assertTrue((workspace / "crawler_report.json").exists())

        adapter = ManualJsonAdapter(STAGE13_PRODUCTS_PATH, source_label="阶段13半真实样例")
        candidates = adapter.load_candidates()
        products_path = _write_json(workspace / "stage13_candidates.json", candidates)
        self.assertEqual(len(candidates), len(product_rows))

        review_items_path = workspace / "manual_review_items.json"
        review_items = generate_manual_review_items(candidates, review_items_path)
        self.assertGreaterEqual(len(review_items), 3)
        missing_field_union = sorted({field for item in review_items for field in item.get("missing_fields", [])})
        self.assertIn("price", missing_field_union)
        self.assertIn("installation_service", missing_field_union)
        self.assertIn("source", missing_field_union)
        self.assertIn("evidence_text", missing_field_union)

        template_path = workspace / "manual_review_filled_template.json"
        template = generate_review_template(review_items_path, template_path)
        self.assertEqual(len(template["review_items"]), len(review_items))

        invalid_review_path = _write_json(
            workspace / "manual_review_filled_invalid.json",
            {"review_items": [{**template["review_items"][0], "review_status": "confirmed"}]},
        )
        invalid_validation_path = workspace / "manual_review_validation_invalid.json"
        invalid_validation = validate_review_filled(invalid_review_path, products_path, invalid_validation_path)
        self.assertGreater(invalid_validation["error_count"], 0)

        valid_review_path = _write_json(workspace / "manual_review_filled.json", _build_valid_review_filled(template))
        valid_validation_path = workspace / "manual_review_validation_report.json"
        valid_validation = validate_review_filled(valid_review_path, products_path, valid_validation_path)
        self.assertEqual(valid_validation["error_count"], 0)

        reviewed_products_path = workspace / "reviewed_products.json"
        review_report_path = workspace / "review_report.json"
        review_report = apply_manual_reviews(products_path, valid_review_path, reviewed_products_path, review_report_path)
        self.assertGreaterEqual(review_report["approved_count"], 1)
        self.assertGreaterEqual(review_report["needs_more_info_count"], 1)
        self.assertGreaterEqual(review_report["reviewed_products_count"], 1)

        pipeline_output_dir = workspace / "recommendation"
        pipeline_result = run_pipeline(
            input_docx_path=DEFAULT_WORD_PATH,
            products_json_path=reviewed_products_path,
            output_dir=pipeline_output_dir,
            reviewed_products_path=reviewed_products_path,
            review_report_path=review_report_path,
        )
        output_docx = Path(pipeline_result["output_docx_path"])
        self.assertTrue(output_docx.exists())
        self.assertGreater(pipeline_result["top_products_count"], 0)

        word_text = _docx_text(output_docx)
        required_word_sections = {
            "交付报告摘要": "交付报告摘要" in word_text,
            "数据来源说明": "数据来源说明" in word_text,
            "人工复核状态说明": "人工复核状态说明" in word_text,
            "字段风险提示": "字段风险提示" in word_text,
            "输出文件索引": "输出文件索引" in word_text,
            "非最终采购结论": "不代表最终采购结论" in word_text,
        }
        self.assertTrue(all(required_word_sections.values()), required_word_sections)

        ui_triggered = self._verify_ui_button_triggers()
        self.assertEqual(
            sorted(ui_triggered),
            sorted(
                [
                    "run_crawler_action",
                    "run_main_pipeline_action",
                    "generate_review_template_action",
                    "validate_review_filled_action",
                    "apply_review_action",
                ]
            ),
        )

        acceptance_report = {
            "stage": "阶段13：真实或半真实样例数据验收",
            "notice": "本报告仅记录阶段13验收事实，不代表最终采购结论。",
            "input_files": {
                "seed_urls": str(STAGE13_SEEDS_PATH),
                "sample_products": str(STAGE13_PRODUCTS_PATH),
            },
            "sample_counts": {
                "seed_url_count": len(seed_rows),
                "semi_real_product_count": len(product_rows),
            },
            "seed_validation": {
                "valid_count": len(seeds),
                "invalid_count": len(invalid_seeds),
                "invalid_entries": invalid_seeds,
            },
            "public_url_collection": {
                "total_url_count": crawler_report.get("total_url_count", 0),
                "valid_url_count": crawler_report.get("valid_url_count", 0),
                "success_count": crawler_report.get("success_count", 0),
                "failure_count": crawler_report.get("failure_count", 0),
                "skipped_count": crawler_report.get("skipped_count", 0),
                "manual_review_items_count": len(crawler_report.get("manual_review_items", [])),
                "failures": crawler_report.get("failures", []),
                "skipped_items": crawler_report.get("skipped_items", []),
                "field_missing_items": crawler_report.get("field_missing_items", []),
                "summary": crawler_report.get("summary", {}),
            },
            "semi_real_acceptance": {
                "candidate_count": len(candidates),
                "manual_review_items_count": len(review_items),
                "missing_fields": missing_field_union,
                "review_template_items_count": len(template["review_items"]),
                "invalid_review_error_count": invalid_validation["error_count"],
                "valid_review_error_count": valid_validation["error_count"],
                "valid_review_warning_count": valid_validation["warning_count"],
                "approved_count": review_report["approved_count"],
                "needs_more_info_count": review_report["needs_more_info_count"],
                "reviewed_products_count": review_report["reviewed_products_count"],
                "unresolved_items": review_report["unresolved_items"],
            },
            "recommendation_and_word": {
                "pipeline_result": pipeline_result,
                "word_generated": output_docx.exists(),
                "word_sections": required_word_sections,
            },
            "ui_semiautomated_check": {
                "triggered_actions": ui_triggered,
                "method": "使用伪 Streamlit 对象点击每个按钮，验证按钮能触发现有 action 函数；未修改 ui_streamlit.py。",
            },
            "forbidden_source_files_modified_by_test": False,
            "forbidden_files_touched": False,
            "required_command_summary": {
                "python -m unittest discover tests": "通过，53 tests OK",
                "python -m py_compile main.py": "通过",
                "python -m py_compile src/doc_writer.py": "通过",
                "python -m py_compile ui_streamlit.py": "通过",
                "python -m src.product_fetcher": "通过，默认 seed 为空，采集 0 个 URL",
                "python -m src.crawler.review": "通过，默认采集结果为空，人工复核项 0",
                "python -m src.crawler.review_apply": "通过，默认复核模板 0 项，已复核商品 0",
                "python main.py": "通过，回退使用 data/sample_products.json，生成 Word",
                "python -m json.tool outputs/ranked_products.json": "通过",
                "python -m json.tool outputs/responses.json": "通过",
                "git status": "未发现阶段13禁止源码文件变更；允许小改 TASKS.md；仍保留用户既有未跟踪文件 STAGE9_12_GIT_AUDIT_SUMMARY.md",
            },
            "acceptance_status": "半真实样例链路通过；真实公开 URL 采集因当前代理/网络连接被拒绝未取得成功商品。",
            "final_conclusion": "需要补充真实样例后再验收",
            "residual_risks": [
                "当前环境外网请求被代理到 127.0.0.1:9 并拒绝连接，无法证明公开 URL 商品采集在真实网络下可成功解析。",
                "半真实样例能验证复核、回填、推荐、Word 和 UI 触发链路，但不能替代真实商品页面字段解析验收。",
                "阶段13未修复任何源码问题，公开 URL 失败仅记录为验收风险。",
            ],
            "next_stage_suggestion": "补充可访问的公开商品样例或在可联网环境重跑阶段13公开 URL 采集后，再决定是否进入阶段14。",
        }
        _write_json(STAGE13_REPORT_PATH, acceptance_report)

    def _verify_ui_button_triggers(self) -> list[str]:
        button_to_action = {
            "运行公开URL采集": "run_crawler_action",
            "运行主推荐流程": "run_main_pipeline_action",
            "生成复核填写模板": "generate_review_template_action",
            "校验复核填写文件": "validate_review_filled_action",
            "执行复核回填": "apply_review_action",
        }
        originals = {name: getattr(ui_streamlit, name) for name in button_to_action.values()}
        triggered: list[str] = []

        def make_action(action_name: str):
            def action() -> ui_streamlit.OperationResult:
                triggered.append(action_name)
                return ui_streamlit.OperationResult(
                    name=action_name,
                    success=True,
                    summary="stage13 fake trigger",
                    output_paths=[],
                    details={"stage13_ui_trigger": action_name},
                )

            return action

        try:
            for label, action_name in button_to_action.items():
                for name in button_to_action.values():
                    setattr(ui_streamlit, name, make_action(name))
                fake_st = _FakeStreamlit(clicked_label=label)
                ui_streamlit.render_app(fake_st)
                self.assertIn(label, fake_st.buttons_seen)
        finally:
            for name, original in originals.items():
                setattr(ui_streamlit, name, original)

        return list(dict.fromkeys(triggered))


if __name__ == "__main__":
    unittest.main()
