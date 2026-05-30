import importlib
import inspect
import sys
import types
import unittest
from pathlib import Path

from docx import Document

from main import build_marketplace_search_links, load_candidate_products_file, run_stage13_2b_delivery
from src.doc_parser import REQUIRED_HEADERS
from src.doc_writer import write_responses
from src.product_ranker import rank_products
from src.response_builder import select_top_products


ROOT = Path(__file__).resolve().parents[1]
DEMO_JSON = ROOT / "data" / "stage13_2b_candidate_products_demo.json"
DEMO_CSV = ROOT / "data" / "stage13_2b_candidate_products_demo.csv"
DEMO_WORD = ROOT / "data" / "办公屏风01模板样例.docx"


class _StreamlitTrap(types.ModuleType):
    def __init__(self) -> None:
        super().__init__("streamlit")
        self.accessed: list[str] = []

    def __getattr__(self, name: str) -> object:
        self.accessed.append(name)
        raise AssertionError(f"streamlit.{name} should not be accessed during import")


def _workspace(name: str) -> Path:
    path = ROOT / "outputs" / "stage13_2b_test_workspace" / name
    path.mkdir(parents=True, exist_ok=True)
    return path


def _create_input_docx(path: Path) -> None:
    document = Document()
    table = document.add_table(rows=2, cols=len(REQUIRED_HEADERS))
    for index, header in enumerate(REQUIRED_HEADERS):
        table.rows[0].cells[index].text = header
    for index, value in enumerate(["1", "屏风尺寸", "mm", "1600x750x750mm，钢制办公屏风", ""]):
        table.rows[1].cells[index].text = value
    document.save(str(path))


def _extract_docx_text(path: Path) -> str:
    document = Document(str(path))
    parts: list[str] = []
    parts.extend(paragraph.text for paragraph in document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


class Stage132BDeliveryMVPTests(unittest.TestCase):
    def test_demo_candidate_json_can_be_read(self) -> None:
        products = load_candidate_products_file(DEMO_JSON)

        self.assertGreaterEqual(len(products), 4)
        self.assertEqual(products[0]["adapter_source_type"], "manual_json")
        self.assertIn("办公屏风", products[0]["title"])
        self.assertIn("现场安装", products[0]["installation_service"])
        self.assertIn("evidence_text", products[0])

    def test_demo_candidate_csv_can_be_read(self) -> None:
        products = load_candidate_products_file(DEMO_CSV)

        self.assertGreaterEqual(len(products), 4)
        self.assertEqual(products[0]["adapter_source_type"], "manual_csv")
        self.assertEqual(products[0]["price"], 998)
        self.assertIn("1600x750x750", products[0]["dimensions"])

    def test_gbk_csv_from_windows_excel_can_be_read(self) -> None:
        workspace = _workspace("gbk_csv")
        csv_path = workspace / "excel_gbk.csv"
        csv_path.write_text(
            "商品名称,平台,价格,链接,来源,尺寸,材质,安装服务,规格参数,证据文本\n"
            "GBK办公屏风 1600x750x750,人工确认候选,899,https://example.com/gbk,客户CSV,1600x750x750mm,钢制框架,支持现场安装,尺寸：1600x750x750mm；材质：钢制框架。,客户CSV提供的演示证据文本。\n",
            encoding="gbk",
        )

        products = load_candidate_products_file(csv_path)

        self.assertEqual(len(products), 1)
        self.assertEqual(products[0]["title"], "GBK办公屏风 1600x750x750")
        self.assertEqual(products[0]["price"], 899)
        self.assertIn("现场安装", products[0]["installation_service"])

    def test_candidate_fields_enter_existing_ranking_shape(self) -> None:
        products = load_candidate_products_file(DEMO_JSON)
        ranked = rank_products(
            products,
            [{"name": "屏风尺寸和材质", "required_value": "办公屏风 1600x750x750mm 钢制框架 现场安装"}],
        )
        top_products = select_top_products(ranked, top_n=3)

        self.assertEqual(len(top_products), 3)
        first = top_products[0]
        for field in ["title", "platform", "price", "url", "dimensions", "material", "service_text", "image_url"]:
            self.assertIn(field, first)
        self.assertIn("现场安装", first["service_text"])

    def test_missing_fields_are_not_fabricated(self) -> None:
        products = load_candidate_products_file(DEMO_JSON)
        missing_product = next(product for product in products if "缺字段" in product["title"])

        self.assertIn("price", missing_product["missing_fields"])
        self.assertIn("dimensions", missing_product["missing_fields"])
        self.assertIn("material", missing_product["missing_fields"])
        self.assertIn("installation_service", missing_product["missing_fields"])
        self.assertEqual(missing_product["price"], None)
        self.assertEqual(missing_product["dimensions"], "")
        self.assertEqual(missing_product["material"], "")

    def test_missing_fields_render_as_manual_review_in_word(self) -> None:
        workspace = _workspace("missing_fields_word")
        input_docx = workspace / "input.docx"
        output_docx = workspace / "output.docx"
        _create_input_docx(input_docx)

        write_responses(
            input_docx,
            [{"index": "1", "response_value": "需人工复核：候选商品字段不足。"}],
            output_docx,
            summary_products=[
                {
                    "title": "缺字段办公屏风候选",
                    "platform": "人工确认候选",
                    "price": None,
                    "url": "https://example.com/missing",
                    "dimensions": "",
                    "material": "",
                    "service_text": "",
                    "image_url": "",
                    "missing_fields": ["price", "dimensions", "material", "installation_service", "image_url"],
                    "manual_review_required": True,
                }
            ],
        )

        text = _extract_docx_text(output_docx)
        self.assertIn("待人工复核", text)
        self.assertIn("价格待人工复核", text)
        self.assertIn("安装服务待人工复核", text)
        self.assertNotIn("Q235", text)
        self.assertNotIn("免费安装", text)

    def test_marketplace_search_keywords_and_links_are_generated(self) -> None:
        search = build_marketplace_search_links(
            [{"name": "办公屏风", "required_value": "1600x750x750mm 钢制框架 灰白色"}]
        )

        self.assertIn("办公屏风", search["taobao_keyword"])
        self.assertIn("1600x750x750", search["jd_keyword"])
        self.assertTrue(search["taobao_search_url"].startswith("https://s.taobao.com/search?q="))
        self.assertTrue(search["jd_search_url"].startswith("https://search.jd.com/Search?keyword="))
        self.assertIn("后续 V2", search["notice"])

    def test_ui_streamlit_imports_without_running_streamlit(self) -> None:
        previous_streamlit = sys.modules.get("streamlit")
        trap = _StreamlitTrap()
        sys.modules["streamlit"] = trap
        sys.modules.pop("ui_streamlit", None)

        try:
            module = importlib.import_module("ui_streamlit")
        finally:
            if previous_streamlit is None:
                sys.modules.pop("streamlit", None)
            else:
                sys.modules["streamlit"] = previous_streamlit

        self.assertTrue(callable(getattr(module, "render_app", None)))
        self.assertTrue(callable(getattr(module, "run_stage13_2b_delivery_action", None)))
        self.assertEqual(trap.accessed, [])

    def test_write_responses_legacy_interface_still_available(self) -> None:
        self.assertEqual(
            list(inspect.signature(write_responses).parameters)[:3],
            ["input_docx_path", "responses", "output_docx_path"],
        )

        workspace = _workspace("legacy_write_responses")
        input_docx = workspace / "input.docx"
        output_docx = workspace / "output.docx"
        _create_input_docx(input_docx)

        result = write_responses(input_docx, [{"index": "1", "response_value": "旧接口仍可写回。"}], output_docx)

        self.assertEqual(result, output_docx)
        self.assertTrue(output_docx.exists())
        self.assertIn("旧接口仍可写回。", _extract_docx_text(output_docx))

    def test_stage13_2b_delivery_generates_downloadable_word(self) -> None:
        workspace = _workspace("delivery")
        result = run_stage13_2b_delivery(DEMO_WORD, DEMO_JSON, output_dir=workspace)

        self.assertGreater(result["requirements_count"], 0)
        self.assertEqual(result["products_count"], 4)
        self.assertEqual(result["top_products_count"], 3)
        self.assertTrue(Path(result["output_docx_path"]).exists())
        text = _extract_docx_text(Path(result["output_docx_path"]))
        self.assertIn("淘宝/京东寻源辅助", text)
        self.assertIn("图片链接", text)
        self.assertIn("待复核字段", text)

    def test_no_positive_final_procurement_conclusion_phrase(self) -> None:
        checked_files = [ROOT / "main.py", ROOT / "ui_streamlit.py", ROOT / "src" / "doc_writer.py"]
        forbidden_positive_phrases = [
            "本报告为最终采购结论",
            "本报告是最终采购结论",
            "作为最终采购结论",
            "可直接下单",
        ]
        combined_text = "\n".join(path.read_text(encoding="utf-8") for path in checked_files)

        for phrase in forbidden_positive_phrases:
            self.assertNotIn(phrase, combined_text)


if __name__ == "__main__":
    unittest.main()
