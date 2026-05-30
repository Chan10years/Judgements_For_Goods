from pathlib import Path
from typing import Any

from docx import Document

from src.doc_parser import REQUIRED_HEADERS


DELIVERY_OUTPUT_FILE_INDEX = [
    "outputs/recommendation_result.docx",
    "outputs/ranked_products.json",
    "outputs/responses.json",
    "outputs/crawler_products.json",
    "outputs/manual_review_items.json",
    "outputs/manual_review_filled_template.json",
    "outputs/manual_review_validation_report.json",
    "outputs/reviewed_products.json",
    "outputs/review_report.json",
    "outputs/sourcing_score_table.csv",
]
MISSING_FIELD_LABELS = {
    "title": "商品名称",
    "platform": "平台",
    "price": "价格",
    "url": "商品链接",
    "shop": "店铺",
    "source": "来源",
    "dimensions": "尺寸",
    "material": "材质",
    "color": "颜色",
    "installation_service": "安装服务",
    "image_url": "图片链接",
    "specs_text": "规格参数",
    "service_text": "服务说明",
    "evidence_text": "证据文本",
}


def _cell_text(cell) -> str:
    return " ".join(cell.text.split())


def _header_map(table) -> dict[str, int] | None:
    if not table.rows:
        return None

    headers = [_cell_text(cell) for cell in table.rows[0].cells]
    mapping: dict[str, int] = {}

    for required_header in REQUIRED_HEADERS:
        try:
            mapping[required_header] = headers.index(required_header)
        except ValueError:
            return None

    return mapping


def _find_requirements_table(document: Document):
    for table in document.tables:
        mapping = _header_map(table)
        if mapping is not None:
            return table, mapping
    raise ValueError("未找到包含技术参数表头的表格。")


def _format_product_value(value: Any, empty_text: str = "") -> str:
    if value is None:
        return empty_text
    if isinstance(value, list):
        text = "；".join(str(item) for item in value if str(item).strip())
        return text or empty_text
    text = str(value).strip()
    return text if text else empty_text


def _first_product_value(product: dict[str, Any], *fields: str, empty_text: str = "") -> str:
    for field in fields:
        text = _format_product_value(product.get(field), "")
        if text:
            return text
    return empty_text


def _format_missing_fields(product: dict[str, Any]) -> str:
    missing_fields = list(product.get("missing_fields") or [])
    computed_fields = {
        "price": product.get("price"),
        "dimensions": product.get("dimensions"),
        "material": product.get("material"),
        "installation_service": _first_product_value(product, "installation_service", "service_text"),
        "url": product.get("url"),
        "image_url": product.get("image_url"),
        "evidence_text": _first_product_value(product, "evidence_text", "evidence"),
    }
    for field, value in computed_fields.items():
        if not _format_product_value(value, "") and field not in missing_fields:
            missing_fields.append(field)

    labels = [MISSING_FIELD_LABELS.get(field, field) for field in dict.fromkeys(missing_fields) if field]
    if not labels:
        return "无"
    return "；".join(f"{label}待人工复核" for label in labels)


def _format_match_notes(product: dict[str, Any]) -> str:
    reasons = _format_product_value(product.get("reasons"), "")
    reason = _format_product_value(product.get("reason"), "")
    score = _format_product_value(product.get("score"), "")
    parts = []
    if score:
        parts.append(f"匹配分数：{score}")
    if reasons:
        parts.append(reasons)
    elif reason:
        parts.append(reason)
    return "；".join(parts) if parts else "仅作为低置信度候选，需人工复核。"


def _format_risk_notes(product: dict[str, Any]) -> str:
    risks = _format_product_value(product.get("risk_tips"), "") or _format_product_value(product.get("risks"), "")
    if product.get("manual_review_required"):
        risks = f"{risks}；存在待人工复核字段。" if risks else "存在待人工复核字段。"
    return risks or "采购前需人工确认价格、规格、来源和服务。"


def _format_recommendation_explanation(product: dict[str, Any]) -> str:
    explanation = product.get("recommendation_explanation")
    if not isinstance(explanation, dict):
        return _format_match_notes(product)
    parts = [
        explanation.get("why_recommended"),
        "命中指标：" + _format_metadata_value(explanation.get("matched_requirements")),
        "排名说明：" + _format_metadata_value(explanation.get("lower_rank_reason")),
    ]
    exclusion = _format_metadata_value(explanation.get("exclusion_reason"))
    if exclusion != "无":
        parts.append("排除提示：" + exclusion)
    return "；".join(_format_product_value(part) for part in parts if _format_product_value(part))


def _format_metadata_value(value: Any) -> str:
    if value is None:
        return "无"
    if isinstance(value, bool):
        return "是" if value else "否"
    if isinstance(value, (list, tuple)):
        return "；".join(_format_metadata_value(item) for item in value) if value else "无"
    if isinstance(value, dict):
        return "；".join(f"{key}：{_format_metadata_value(item)}" for key, item in value.items()) if value else "无"
    text = str(value).strip()
    return text if text else "无"


def _metadata_int(metadata: dict[str, Any], key: str) -> int:
    value = metadata.get(key, 0)
    try:
        return int(value or 0)
    except (TypeError, ValueError):
        return 0


def _append_key_value_table(document: Document, rows: list[tuple[str, Any]]) -> None:
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "项目"
    table.rows[0].cells[1].text = "内容"

    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = _format_metadata_value(value)


def _add_paragraph(document: Document, text: str, style: str | None = None) -> None:
    if style is None:
        document.add_paragraph(text)
        return
    try:
        document.add_paragraph(text, style=style)
    except KeyError:
        document.add_paragraph(text)


def _append_delivery_summary(document: Document) -> None:
    document.add_heading("交付报告摘要", level=1)
    document.add_paragraph(
        "本报告为采购辅助候选报告，基于当前可用商品数据和本地规则生成候选与响应建议；"
        "所有候选、字段和证据均需要人工复核，不代表最终采购结论。"
    )


def _append_data_source_section(document: Document, metadata: dict[str, Any]) -> None:
    priority = metadata.get(
        "data_source_priority",
        ["reviewed_products.json", "crawler_products.json", "sample_products.json"],
    )

    document.add_heading("数据来源说明", level=1)
    document.add_paragraph("本次商品数据来源优先级如下：")
    for source in priority:
        _add_paragraph(document, _format_metadata_value(source), style="List Number")

    _append_key_value_table(
        document,
        [
            ("当前实际使用数据源", metadata.get("actual_data_source", "未识别")),
            ("当前数据源路径", metadata.get("products_source_path", "未提供")),
            ("说明", metadata.get("data_source_note", "按可用数据源优先级自动选择。")),
        ],
    )


def _append_recommendation_summary_section(document: Document, metadata: dict[str, Any]) -> None:
    output_paths = metadata.get("output_paths", {})
    document.add_heading("推荐结果摘要", level=1)
    _append_key_value_table(
        document,
        [
            ("需求数量", _metadata_int(metadata, "requirements_count")),
            ("候选商品数量", _metadata_int(metadata, "products_count")),
            ("排序后候选商品数量", _metadata_int(metadata, "ranked_products_count")),
            ("Top候选数量", _metadata_int(metadata, "top_products_count")),
            ("Word输出文件路径", output_paths.get("recommendation_result_docx")),
            ("排序结果文件路径", output_paths.get("ranked_products_json")),
            ("响应结果文件路径", output_paths.get("responses_json")),
            ("评分表 CSV 路径", output_paths.get("sourcing_score_table_csv")),
        ],
    )


def _append_review_status_section(document: Document, metadata: dict[str, Any]) -> None:
    review_status = metadata.get("review_status", {})
    if not isinstance(review_status, dict):
        review_status = {}

    document.add_heading("人工复核状态说明", level=1)
    _append_key_value_table(
        document,
        [
            ("存在 reviewed_products.json", review_status.get("has_reviewed_products", False)),
            ("存在 review_report.json", review_status.get("has_review_report", False)),
            ("已复核商品数量", review_status.get("reviewed_products_count", 0)),
            ("approved 数量", review_status.get("approved_count", 0)),
            ("rejected 数量", review_status.get("rejected_count", 0)),
            ("needs_more_info 数量", review_status.get("needs_more_info_count", 0)),
            ("未匹配复核项数量", review_status.get("unmatched_review_items_count", 0)),
            ("仍需人工复核提示", review_status.get("risk_note", "仍需人工确认候选商品与关键字段。")),
        ],
    )


def _append_field_risk_section(document: Document) -> None:
    document.add_heading("字段风险提示", level=1)
    document.add_paragraph(
        "价格、尺寸、材质、安装服务、来源、证据文本等字段如缺失、为空或证据不足，"
        "需要人工确认；未确认字段不得写成已确认，也不得直接作为采购决策依据。"
    )
    for item in [
        "价格：如缺少币种、含税口径、运费或时效信息，需要人工确认。",
        "尺寸：如未覆盖长宽高、厚度、屏风高度或适配场景，需要人工确认。",
        "材质：如只有营销描述、没有明确材料或结构信息，需要人工确认。",
        "安装服务：如未说明是否含安装、配送、售后或现场条件，需要人工确认。",
        "来源与证据文本：如缺少链接、来源、截图摘录或证据文本，需要人工确认。",
    ]:
        _add_paragraph(document, item, style="List Bullet")


def _append_marketplace_search_section(document: Document, metadata: dict[str, Any]) -> None:
    search = metadata.get("marketplace_search", {})
    if not isinstance(search, dict) or not search:
        return

    document.add_heading("淘宝/京东寻源辅助", level=1)
    document.add_paragraph(
        search.get(
            "notice",
            "当前版本提供淘宝/京东寻源辅助入口。候选商品需通过人工确认或候选商品文件导入后进入自动筛选排序和 Word 输出流程。"
            "淘宝/京东全自动搜索、商品详情抓取和图片证据处理属于后续 V2。",
        )
    )
    _append_key_value_table(
        document,
        [
            ("精准搜索词", search.get("precise_terms", [])),
            ("放宽搜索词", search.get("relaxed_terms", [])),
            ("替代搜索词", search.get("alternative_terms", [])),
            ("建议排除词", search.get("excluded_terms", [])),
            ("淘宝搜索关键词", search.get("taobao_keyword", "办公屏风 工位")),
            ("京东搜索关键词", search.get("jd_keyword", "办公屏风 工位")),
            ("淘宝搜索链接", search.get("taobao_search_url", "待人工复核")),
            ("京东搜索链接", search.get("jd_search_url", "待人工复核")),
        ],
    )

    suggestions = metadata.get("platform_filter_suggestions") or search.get("filter_suggestions") or []
    if suggestions:
        document.add_paragraph("平台筛选建议（采购辅助）：")
        for item in suggestions:
            if isinstance(item, dict):
                text = f"{item.get('category', '建议')}：{item.get('suggestion', '')}"
            else:
                text = _format_metadata_value(item)
            _add_paragraph(document, text, style="List Bullet")


def _append_output_index_section(document: Document, metadata: dict[str, Any]) -> None:
    files = metadata.get("output_file_index") or DELIVERY_OUTPUT_FILE_INDEX
    document.add_heading("输出文件索引", level=1)
    for file_path in files:
        _add_paragraph(document, _format_metadata_value(file_path), style="List Bullet")


def _append_delivery_report(document: Document, delivery_metadata: dict[str, Any]) -> None:
    document.add_page_break()
    _append_delivery_summary(document)
    _append_data_source_section(document, delivery_metadata)
    _append_recommendation_summary_section(document, delivery_metadata)
    _append_review_status_section(document, delivery_metadata)
    _append_field_risk_section(document)
    _append_marketplace_search_section(document, delivery_metadata)
    _append_output_index_section(document, delivery_metadata)


def _append_summary_table(document: Document, summary_products: list[dict[str, Any]]) -> None:
    document.add_paragraph()
    document.add_paragraph("候选商品 Top 3（本地规则排序候选商品，非最终采购结论）")

    headers = [
        "排名",
        "商品名称",
        "平台",
        "价格",
        "尺寸",
        "材质",
        "安装服务",
        "商品链接",
        "图片链接",
        "匹配说明",
        "待复核字段",
        "风险提示",
        "推荐等级",
        "推荐解释",
        "人工确认问题",
    ]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header

    for rank, product in enumerate(summary_products, start=1):
        row = table.add_row().cells
        row[0].text = str(rank)
        row[1].text = _format_product_value(product.get("title"), "待人工复核")
        row[2].text = _first_product_value(product, "platform", "source", empty_text="待人工复核")
        row[3].text = _format_product_value(product.get("price"), "待人工复核")
        row[4].text = _format_product_value(product.get("dimensions"), "待人工复核")
        row[5].text = _format_product_value(product.get("material"), "待人工复核")
        row[6].text = _first_product_value(product, "installation_service", "service_text", empty_text="待人工复核")
        row[7].text = _format_product_value(product.get("url"), "待人工复核")
        row[8].text = _format_product_value(product.get("image_url"), "待人工复核")
        row[9].text = _format_match_notes(product)
        row[10].text = _format_missing_fields(product)
        row[11].text = _format_risk_notes(product)
        row[12].text = _format_product_value(product.get("recommendation_level"), "待人工复核")
        row[13].text = _format_recommendation_explanation(product)
        row[14].text = _format_product_value(product.get("manual_confirmation_questions"), "待人工复核")


def write_responses(
    input_docx_path: str | Path,
    responses: list[dict[str, str]],
    output_docx_path: str | Path,
    summary_products: list[dict[str, Any]] | None = None,
    delivery_metadata: dict[str, Any] | None = None,
) -> Path:
    document = Document(str(input_docx_path))
    table, mapping = _find_requirements_table(document)
    response_by_index = {item.get("index", ""): item.get("response_value", "") for item in responses}

    for row_number, row in enumerate(table.rows[1:]):
        cells = row.cells
        index = _cell_text(cells[mapping["序号"]])
        response_value = response_by_index.get(index)
        if response_value is None and row_number < len(responses):
            response_value = responses[row_number].get("response_value", "")
        cells[mapping["投标人响应值"]].text = response_value or ""

    if summary_products:
        _append_summary_table(document, summary_products)

    if delivery_metadata:
        _append_delivery_report(document, delivery_metadata)

    output_path = Path(output_docx_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    return output_path
